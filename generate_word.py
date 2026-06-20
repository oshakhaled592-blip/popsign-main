from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.18)
    section.right_margin  = Cm(3.18)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, name='Times New Roman', size=12, bold=False,
             italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level):
    """H2 → section, H3 → subsection, H4 → sub-subsection."""
    sizes  = {1: 20, 2: 16, 3: 14, 4: 13}
    bolds  = {1: True, 2: True, 3: True, 4: True}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    set_font(r, size=sizes.get(level, 12), bold=bolds.get(level, False))
    if level == 2:
        r.font.underline = True
    return p

def add_para(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    set_font(r)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(4)
    # parse bold inline (**text**)
    parts = text.split('**')
    for i, part in enumerate(parts):
        if part == '':
            continue
        r = p.add_run(part)
        set_font(r, bold=(i % 2 == 1))
    return p

def add_code(lines):
    """Grey monospace block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    # shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F0F0F0')
    pPr.append(shd)
    r = p.add_run('\n'.join(lines))
    set_font(r, name='Courier New', size=9)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        set_font(r, size=11, bold=True, color=(255,255,255))
        # header background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '2D2D44')
        tcPr.append(shd)
    # data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri+1]
        fill = 'F7F7F7' if ri % 2 == 1 else 'FFFFFF'
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(cell_text)
            set_font(r, size=11)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  fill)
            tcPr.append(shd)
    doc.add_paragraph()  # spacing after table

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(8):
    doc.add_paragraph()

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cover.add_run('Chapter 7 (Flutter)')
set_font(cr, size=28, bold=True)

for _ in range(20):
    doc.add_paragraph()

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7.1  Introduction to Flutter and Dart
# ══════════════════════════════════════════════════════════════════════════════
add_heading('7.1  Introduction to Flutter and Dart', 2)

add_heading('1. What is Flutter?', 3)
add_para('Flutter is an open-source UI software development toolkit created by Google. It is used for building natively compiled applications for mobile (Android and iOS), web, desktop, and even embedded devices, all from a single codebase. Launched officially in 2017, Flutter quickly gained popularity due to its performance, flexibility, and developer-friendly architecture.')

add_heading('Key Features of Flutter:', 4)
add_bullet('**Cross-Platform Development:** Write your code once and run it anywhere. No need to maintain separate codebases for Android, iOS, Web, or Desktop.')
add_bullet('**Native Performance:** Flutter compiles down to native ARM code using Dart, which ensures smooth performance, unlike other hybrid frameworks that rely on JavaScript bridges.')
add_bullet('**Beautiful UI:** Flutter provides a rich set of widgets that mimic the behavior of native components. Developers can easily build beautiful and custom user interfaces.')
add_bullet('**Hot Reload:** This allows developers to instantly see the changes in the UI without restarting the app, significantly speeding up the development process.')
add_bullet('**Custom Rendering Engine:** Flutter uses its own rendering engine (Skia/Impeller) which means the UI looks and behaves consistently across all platforms.')

add_heading('2. Understanding Dart: The Language Behind Flutter', 3)
add_para('Dart is the programming language used to build Flutter applications. It is also developed by Google and was designed specifically to create fast, modern, reactive apps.')

add_heading('Why Dart?', 4)
add_bullet('**Optimized for UI:** Dart is built to handle UI development. It allows for layout expressions, smooth animations, and declarative programming.')
add_bullet('**Object-Oriented and Class-Based:** Dart follows modern programming principles with strong support for classes, inheritance, and interfaces.')
add_bullet('**Ahead-of-Time (AOT) & Just-in-Time (JIT) Compilation:** AOT compiles the code for fast startup and better performance in production. JIT allows for hot reload during development.')
add_bullet('**Strongly Typed with Null Safety:** Prevents null reference errors at compile time, improving application stability.')
add_bullet('**Garbage Collected and Memory Safe:** Helps manage resources efficiently and prevents many common memory issues.')

add_heading('3. The Relationship Between Flutter and Dart', 3)
add_para('Flutter and Dart are designed to work hand-in-hand. Dart powers the logic and behavior, while Flutter provides the UI framework built entirely with Dart.')
add_bullet('**Single Language Stack:** Unlike other frameworks (like React Native, which uses JavaScript with some native code), Flutter allows the entire application to be written in Dart, simplifying development and debugging.')
add_bullet('**No Need for OEM Widgets:** Flutter draws UI components itself using Dart and Skia, rather than relying on the platform\'s native components. This ensures uniform appearance and behavior.')
add_bullet('**Customizability:** Because everything in Flutter is a widget, developers have full control over appearance and functionality.')

add_heading('4. Flutter Architecture Overview', 3)
add_para('Flutter follows a layered architecture consisting of:')
add_bullet('**Framework Layer:** The top layer, written in Dart, includes core libraries, rendering system, widgets, and foundation classes.')
add_bullet('**Engine Layer:** Written in C++, the engine uses Skia/Impeller to render everything and handles input, text layout, and more.')
add_bullet('**Embedder Layer:** Platform-specific code that allows Flutter to run on various platforms like Android, iOS, web, etc.')

add_heading('5. Advantages of Choosing Flutter for I Hear You', 3)
add_para('The I Hear You application targets sign language learners on both Android and iOS. Flutter was selected because:')
add_bullet('A single codebase covers both platforms simultaneously, reducing development time and cost.')
add_bullet('The rich widget ecosystem supports complex UI interactions like flashcard animations and video playback.')
add_bullet('The flutter_screenutil package ensures fully responsive layouts across all screen sizes.')
add_bullet('Firebase integration is first-class with official Flutter packages.')
add_bullet('Hot Reload enabled fast iteration during development.')

add_heading('6. Challenges and Considerations', 3)
add_bullet('**App Size:** Flutter apps can be slightly larger in size due to the built-in rendering engine.')
add_bullet('**Limited Native SDK Access:** Although Flutter supports platform channels, working with native APIs still requires writing native platform-specific code in Kotlin/Swift.')
add_bullet('**Learning Curve:** Dart is not as widely known as JavaScript or Kotlin, requiring an initial learning investment.')

add_heading('7. Summary', 3)
add_para('Flutter and Dart together provide a robust ecosystem for building modern, responsive, and high-performance applications across multiple platforms. Their integration creates a development experience that is both productive and powerful, with the ability to produce beautiful UIs and maintain a single codebase. This is precisely why Flutter was chosen for the I Hear You sign-language learning application.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7.2  Project Structure and State Management
# ══════════════════════════════════════════════════════════════════════════════
add_heading('7.2  Project Structure and State Management\n(ValueNotifier and ChangeNotifier)', 2)

add_heading('1. Introduction to Project Structure in Flutter', 3)
add_para('Building a scalable and maintainable Flutter application starts with a well-organized project structure. Without proper organization, even small apps can become difficult to manage over time. Flutter doesn\'t impose a specific architecture by default, so developers choose patterns that best fit their project\'s complexity.')
add_para('For the I Hear You application, a feature-based layered architecture was adopted, separating concerns clearly into core utilities and feature modules. This enables scalability without the overhead of complex global state management libraries.')

add_heading('2. Project Folder Structure', 3)
add_code([
    'lib/',
    '├── core/',
    '│   ├── helpers/',
    '│   │   └── spacing.dart',
    '│   ├── l10n/',
    '│   │   └── app_strings.dart',
    '│   ├── models/',
    '│   │   └── prediction_model.dart',
    '│   ├── routing/',
    '│   │   ├── app_router.dart',
    '│   │   └── routes.dart',
    '│   ├── services/',
    '│   │   ├── auth_service.dart',
    '│   │   ├── chat_service.dart',
    '│   │   └── translation_service.dart',
    '│   ├── state/',
    '│   │   └── translation_notifier.dart',
    '│   ├── theme/',
    '│   │   ├── app_colors.dart',
    '│   │   ├── language_notifier.dart',
    '│   │   ├── styles.dart',
    '│   │   └── theme_notifier.dart',
    '│   └── widgets/',
    '│       ├── linear_button.dart',
    '│       ├── logo_and_name.dart',
    '│       └── text_form_widget.dart',
    '├── features/',
    '│   ├── auth/',
    '│   │   ├── login/ui/screen/',
    '│   │   │   ├── login_screen.dart',
    '│   │   │   └── forgot_password_screen.dart',
    '│   │   └── signup/ui/screen/',
    '│   │       ├── signup_screen.dart',
    '│   │       └── check_email_screen.dart',
    '│   └── initialization/ui/',
    '│       ├── screens/',
    '│       │   ├── splash_screen.dart',
    '│       │   ├── dashboard_screen.dart',
    '│       │   ├── levels_screen.dart',
    '│       │   ├── word_list_screen.dart',
    '│       │   ├── sign_practice_screen.dart',
    '│       │   ├── chat_screen.dart',
    '│       │   ├── lost_page_screen.dart',
    '│       │   └── profile._screen.dart',
    '│       └── widgets/',
    '│           └── quiz_step_scaffold.dart',
    '└── main.dart',
])

add_heading('Layer Descriptions:', 4)
add_bullet('**core/helpers/:** Utility functions such as verticalSpace() and horizontalSpace() for consistent widget spacing.')
add_bullet('**core/l10n/:** All user-facing strings across 14 languages via app_strings.dart.')
add_bullet('**core/models/:** Data models for API responses (e.g., PredictionResult, PredictionItem).')
add_bullet('**core/routing/:** Named route definitions and AppRouter for centralised navigation.')
add_bullet('**core/services/:** Service classes abstracting all external communication.')
add_bullet('**core/state/:** TranslationNotifier — a ChangeNotifier managing translation feature state.')
add_bullet('**core/theme/:** Theme configuration, text styles, colour definitions, and notifiers for dark/light mode and language.')
add_bullet('**core/widgets/:** Reusable components: LinearButton, TextFormWidget, and LogoAndName.')

add_heading('3. What is State Management in Flutter?', 3)
add_para('In Flutter, everything is a widget. These widgets can be stateless (don\'t change) or stateful (change over time). Managing the changes in these widgets — especially when the UI depends on data from API responses or user actions — is referred to as State Management.')
add_para('Flutter offers several approaches: setState for local widget state, ValueNotifier for lightweight global reactive state, ChangeNotifier for feature-specific complex state, SharedPreferences for persistent state across sessions, and larger libraries like Provider, Bloc/Cubit, Riverpod, or GetX for enterprise-scale apps.')
add_para('For the I Hear You application, a lightweight, pragmatic approach was chosen — using ValueNotifier, ChangeNotifier, SharedPreferences, and setState — without heavy global state libraries, keeping the codebase lean and maintainable.')

add_heading('4. Understanding ValueNotifier and State', 3)
add_para('A ValueNotifier<T> is a class that holds a single value and notifies all registered listeners whenever that value changes. It is Flutter\'s simplest built-in reactive primitive.')

add_heading('4.1  Theme Notifier — Dark / Light Mode', 4)
add_para('A global ValueNotifier<ThemeMode> manages dark/light mode for the entire app:')
add_code([
    '// theme_notifier.dart',
    'final themeNotifier = ValueNotifier<ThemeMode>(ThemeMode.dark);',
])
add_para('Used in main.dart to rebuild the entire MaterialApp when the user toggles the theme:')
add_code([
    'ValueListenableBuilder<ThemeMode>(',
    '  valueListenable: themeNotifier,',
    '  builder: (context, mode, _) => MaterialApp(',
    '    themeMode: mode,',
    '    theme: ThemeData.light(),',
    '    darkTheme: ThemeData.dark(),',
    '    onGenerateRoute: AppRouter.generateRoute,',
    '  ),',
    ')',
])
add_para('The theme choice is persisted to SharedPreferences (key: isDark) and loaded before runApp() in main.dart so the splash screen always shows the correct theme immediately — eliminating the flash-of-wrong-theme problem:')
add_code([
    'void main() async {',
    '  WidgetsFlutterBinding.ensureInitialized();',
    '  await Firebase.initializeApp(options: DefaultFirebaseOptions.currentPlatform);',
    '  final prefs = await SharedPreferences.getInstance();',
    '  final savedDark = prefs.getBool(\'isDark\');',
    '  if (savedDark != null) {',
    '    themeNotifier.value = savedDark ? ThemeMode.dark : ThemeMode.light;',
    '  }',
    '  runApp(const IHeartYouApp());',
    '}',
])

add_heading('4.2  Language Notifier', 4)
add_para('A ValueNotifier<LanguageInfo> holds the currently selected UI language (code, name, and flag):')
add_code([
    '// language_notifier.dart',
    "final languageNotifier = ValueNotifier<LanguageInfo>(",
    "  LanguageInfo(code: 'en', name: 'English', flag: '🇺🇸'),",
    ');',
])

add_heading('5. ChangeNotifier — Translation Feature State', 3)
add_para('For the sign-language translation feature, a TranslationNotifier extending ChangeNotifier manages the full lifecycle of a video prediction request:')
add_code([
    'enum TranslationStatus { idle, loading, success, error }',
    'enum ServerStatus { unknown, online, offline }',
    '',
    'class TranslationNotifier extends ChangeNotifier {',
    '  final _service = TranslationService();',
    '  bool _disposed = false;',
    '',
    '  @override',
    '  void dispose() {',
    '    _disposed = true;',
    '    super.dispose();',
    '  }',
    '',
    '  void _notify() {',
    '    if (!_disposed) notifyListeners();',
    '  }',
    '',
    '  TranslationStatus _status = TranslationStatus.idle;',
    '  PredictionResult? _result;',
    '  String _error = \'\';',
    '',
    '  Future<void> predict(File video) async {',
    '    _status = TranslationStatus.loading;',
    '    _notify();',
    '    try {',
    '      _result = await _service.predict(video);',
    '      _status = TranslationStatus.success;',
    '    } catch (e) {',
    '      _error = e.toString().replaceFirst(\'Exception: \', \'\');',
    '      _status = TranslationStatus.error;',
    '    }',
    '    _notify();',
    '  }',
    '}',
])
add_para('The _disposed flag is a critical safety pattern: it prevents notifyListeners() from being called after the widget is removed from the tree — protecting against the ChangeNotifier.debugAssertNotDisposed crash when async operations outlive their widget.')

add_heading('6. SharedPreferences — Persistent State', 3)
add_para('All data that must survive app restarts is stored in SharedPreferences:')
add_table(
    ['Key', 'Type', 'Purpose'],
    [
        ['isDark',           'bool',   'Theme preference (dark/light)'],
        ['isLoggedIn',       'bool',   'Authentication gate at startup'],
        ['hasSelectedLevel', 'bool',   'Skip onboarding on next launch'],
        ['userLevel',        'String', 'Selected CEFR level (A1–C2)'],
        ['authToken',        'String', 'Backend JWT token'],
        ['userEmail',        'String', 'Logged-in user email'],
        ['userName',         'String', 'Display name (editable)'],
        ['langCode',         'String', 'Selected learning language code'],
        ['langName',         'String', 'Selected learning language name'],
        ['langFlag',         'String', 'Flag emoji for selected language'],
        ['userPurpose',      'String', 'Reason for learning (family, work…)'],
        ['userMinutesPerDay','int',    'Daily study time goal'],
    ]
)

add_heading('7. setState — Local Screen State', 3)
add_para('Individual screens use setState for transient UI state such as form input, loading indicators, and selection state. Every async method follows the mounted guard pattern:')
add_code([
    'bool _isLoading = false;',
    '',
    'Future<void> _login() async {',
    '  if (!_formKey.currentState!.validate()) return;',
    '  setState(() => _isLoading = true);',
    '',
    '  final result = await AuthService().login(',
    '    email: _emailCtrl.text.trim(),',
    '    password: _passCtrl.text.trim(),',
    '  );',
    '',
    '  if (!mounted) return;  // guard against disposed widget',
    '  setState(() => _isLoading = false);',
    '}',
])

add_heading('8. Summary', 3)
add_para('The I Hear You app uses a four-layer state management model: ValueNotifier for global reactive app state (theme and language), ChangeNotifier for feature-specific complex state (translation), SharedPreferences for persistent cross-session data, and setState for transient screen-level state. This combination provides reactivity where needed without adding unnecessary overhead.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7.3  UI Components and Navigation
# ══════════════════════════════════════════════════════════════════════════════
add_heading('7.3  UI Components and Navigation in Flutter', 2)

add_heading('1. Introduction', 3)
add_para('Flutter is a UI toolkit designed to enable fast and expressive development of cross-platform applications. One of Flutter\'s core strengths lies in its rich set of UI components (Widgets) and its powerful, customizable navigation system. In the I Hear You application, both are used extensively to create a polished, multilingual, and theme-aware experience.')

add_heading('2. Understanding UI in Flutter: Everything is a Widget', 3)
add_para('Flutter takes a widget-centric approach. Everything in Flutter — from layout elements to text and icons — is a widget. There are two main types:')
add_bullet('**Stateless Widgets:** These are immutable. Once created, they do not change. Used when the UI depends only on input values.')
add_bullet('**Stateful Widgets:** These are mutable and can rebuild themselves when their state changes. Useful for dynamic interfaces.')

add_heading('3. Design System', 3)
add_para('The app uses a consistent design system built around two themes — Dark and Light — toggled in real time via the themeNotifier.')

add_heading('3.1  Colour Palette', 4)
add_table(
    ['Role', 'Dark Mode', 'Light Mode'],
    [
        ['Background',    '#0B1020 / #0D0F1A',           '#F5F7FB'],
        ['Card Surface',  '#141829 / #181830',            '#FFFFFF'],
        ['Primary Accent','#22C55E (green)',              '#22C55E (green)'],
        ['Text Primary',  '#FFFFFF',                      '#1A1A2E'],
        ['Text Secondary','rgba(255,255,255,0.54)',        '#6B7280'],
    ]
)

add_heading('3.2  Typography (AppTextStyles)', 4)
add_para('All font sizes use .sp from flutter_screenutil for proper scaling across devices:')
add_code([
    '// styles.dart',
    'static TextStyle font36BoldWhite = GoogleFonts.openSans(',
    '  fontWeight: FontWeightHelper.bold,',
    '  fontSize: 36.sp,',
    '  color: Colors.white,',
    ');',
    '',
    'static TextStyle font16MediumWhite = GoogleFonts.openSans(',
    '  fontWeight: FontWeightHelper.medium,',
    '  fontSize: 16.sp,',
    '  color: Colors.white,',
    ');',
])

add_heading('4. Responsive Design with flutter_screenutil', 3)
add_para('Every dimension in the app uses flutter_screenutil with a design size of 390×844 (iPhone 14 equivalent). This ensures the UI scales correctly across all Android and iOS screen sizes without hardcoded pixel values:')
add_code([
    '// main.dart',
    'ScreenUtilInit(',
    '  designSize: const Size(390, 844),',
    '  minTextAdapt: true,',
    '  builder: (context, child) {',
    '    return MaterialApp( /* ... */ );',
    '  },',
    ')',
])
add_table(
    ['Suffix', 'Scales', 'Example'],
    [
        ['.w',  'Width proportionally',          '24.w'],
        ['.h',  'Height proportionally',         '16.h'],
        ['.sp', 'Font size with text scale',     '14.sp'],
        ['.r',  'Border radius / uniform size',  '12.r'],
    ]
)

add_heading('5. Reusable Custom Widgets', 3)
add_para('To avoid repetition and improve code readability, common UI components are extracted into custom widgets in the core/widgets/ directory.')

add_heading('5.1  LinearButton — Primary CTA Button', 4)
add_para('A gradient action button used across all CTA screens:')
add_code([
    "LinearButton(",
    "  text: S.get('get_started'),",
    "  onPressed: () => Navigator.pushNamed(context, Routes.login),",
    "  width: double.infinity.w,",
    "  radius: 14.r,",
    ")",
])

add_heading('5.2  TextFormWidget — Styled Input Field', 4)
add_code([
    'TextFormWidget(',
    '  controller: _emailCtrl,',
    "  hintText: S.get('email'),",
    '  icon: Icons.email_outlined,',
    '  keyboardType: TextInputType.emailAddress,',
    '  validator: (v) {',
    "    final val = (v ?? '').trim();",
    "    if (val.isEmpty) return S.get('email_required');",
    "    if (!val.contains('@') || !val.contains('.')) {",
    "      return S.get('email_invalid');",
    '    }',
    '    return null;',
    '  },',
    ')',
])

add_heading('5.3  LogoAndName — Theme-Aware App Logo', 4)
add_code([
    'class LogoAndName extends StatelessWidget {',
    '  @override',
    '  Widget build(BuildContext context) {',
    '    final isDark = Theme.of(context).brightness == Brightness.dark;',
    '    return Column(',
    '      children: [',
    "        SvgPicture.asset('assets/svgs/app-logo.svg', height: 81.h, width: 67.w),",
    '        Text(',
    '          "I Hear You",',
    '          style: AppTextStyles.font36SemiboldWhite.copyWith(',
    '            color: isDark ? Colors.white : const Color(0xFF1A1A2E),',
    '          ),',
    '        ),',
    '      ],',
    '    );',
    '  }',
    '}',
])

add_heading('6. Navigation', 3)
add_para('The app uses Named Routes with a central AppRouter class that maps route strings to screen widgets:')
add_code([
    '// routes.dart',
    'class Routes {',
    "  static const String splash    = '/';",
    "  static const String login     = '/login';",
    "  static const String signup    = '/signup';",
    "  static const String dashboard = '/dashboard';",
    "  static const String levels    = '/levels';",
    "  static const String chat      = '/chat';",
    "  static const String profile   = '/profile';",
    '}',
    '',
    '// app_router.dart',
    'class AppRouter {',
    '  static Route<dynamic> generateRoute(RouteSettings settings) {',
    '    switch (settings.name) {',
    '      case Routes.splash:',
    '        return MaterialPageRoute(builder: (_) => const SplashScreen());',
    '      case Routes.login:',
    '        return MaterialPageRoute(builder: (_) => const LoginScreen());',
    '      case Routes.dashboard:',
    '        return MaterialPageRoute(builder: (_) => const DashboardScreen());',
    '      default:',
    '        return MaterialPageRoute(builder: (_) => const SplashScreen());',
    '    }',
    '  }',
    '}',
])

add_heading('6.1  Navigation Flow', 4)
add_para('The SplashScreen reads SharedPreferences flags (isLoggedIn, hasSelectedLevel) and routes the user to the correct starting point:')
add_code([
    '// splash_screen.dart',
    'Future<void> _navigate() async {',
    '  final prefs = await SharedPreferences.getInstance();',
    "  final isLoggedIn       = prefs.getBool('isLoggedIn') ?? false;",
    "  final hasSelectedLevel = prefs.getBool('hasSelectedLevel') ?? false;",
    '  if (!mounted) return;',
    '  if (isLoggedIn && hasSelectedLevel) {',
    '    Navigator.pushReplacementNamed(context, Routes.dashboard);',
    '  } else if (isLoggedIn) {',
    '    Navigator.pushReplacementNamed(context, Routes.chooseLanguage);',
    '  } else {',
    '    Navigator.pushReplacementNamed(context, Routes.startPage);',
    '  }',
    '}',
])

add_heading('7. Commonly Used UI Widgets', 3)
add_table(
    ['Widget', 'Description'],
    [
        ['Scaffold',                  'Base layout for all screens with SafeArea wrapping'],
        ['SingleChildScrollView',     'Prevents overflow on small screens with keyboards'],
        ['ListView.builder',          'Efficient scrolling lists in LevelsScreen and WordListScreen'],
        ['AnimatedContainer',         'Smooth transitions in flashcard and progress animations'],
        ['CircularProgressIndicator', 'Loading states across all API calls'],
        ['SnackBar',                  'User feedback messages (error, success, "Coming Soon")'],
        ['AlertDialog',               'Success confirmation dialogs (name edit, email sent)'],
        ['Switch',                    'Dark mode toggle in ProfileScreen'],
        ['VideoPlayer',               'Sign practice video preview via video_player package'],
    ]
)

add_heading('8. Summary', 3)
add_para('The I Hear You app demonstrates Flutter\'s UI capabilities through a consistent design system, a library of reusable custom widgets, and a centralized named-route navigation system. Every dimension uses flutter_screenutil for pixel-perfect responsiveness, and every text colour is derived from the active theme — ensuring a polished experience in both Dark and Light modes.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7.4  API Integration using HTTP
# ══════════════════════════════════════════════════════════════════════════════
add_heading('7.4  API Integration using HTTP in Flutter', 2)

add_heading('1. Introduction', 3)
add_para('Modern mobile apps rarely operate in isolation; they communicate with remote servers through RESTful APIs to fetch or send data. In Flutter, this is typically done using the http package — the standard Dart HTTP client — or the more advanced dio package.')
add_para('The I Hear You application integrates three external servers:')
add_bullet('A **backend REST API** at http://91.108.113.135 — primary authentication source')
add_bullet('A **Firebase Auth** server — for password management operations')
add_bullet('A **dedicated AI server** — for sign-language prediction and conversational chat')
add_para('All external communication is abstracted into service classes (AuthService, TranslationService, ChatService), keeping UI code free of HTTP concerns.')

add_heading('2. Using the http Package', 3)
add_para('The http package is Flutter\'s standard HTTP client for making REST API calls.')

add_heading('2.1  Installation', 4)
add_code([
    'dependencies:',
    '  http: ^1.2.1',
])

add_heading('2.2  The AuthService Pattern', 4)
add_para('Every method in AuthService returns an AuthResult with a success flag and a user-facing message, keeping error handling uniform across the app:')
add_code([
    'class AuthResult {',
    '  final bool success;',
    '  final String message;',
    '  AuthResult({required this.success, required this.message});',
    '}',
    '',
    'class AuthService {',
    "  static const _base = 'http://91.108.113.135';",
    '  static const _timeout = Duration(seconds: 20);',
    '  final _fb = FirebaseAuth.instance;',
    '}',
])

add_heading('2.3  Making a POST Request — Login', 4)
add_code([
    'Future<AuthResult> login({',
    '  required String email,',
    '  required String password,',
    '}) async {',
    '  try {',
    '    final res = await http.post(',
    "      Uri.parse('$_base/api/auth/login'),",
    "      headers: {'Content-Type': 'application/json'},",
    "      body: jsonEncode({'email': email, 'password': password}),",
    '    ).timeout(_timeout);',
    '',
    '    final json = _decode(res);',
    '    final token = _extractToken(json);',
    '    if (token != null) {',
    '      final prefs = await SharedPreferences.getInstance();',
    "      await prefs.setString('authToken', token);",
    "      await prefs.setString('userEmail', email);",
    '    }',
    '',
    '    if (json[\'success\'] != true) {',
    '      return AuthResult(',
    '        success: false,',
    "        message: json['message']?.toString() ?? 'Login failed',",
    '      );',
    '    }',
    '',
    '    await _signIntoFirebase(email, password);',
    "    return AuthResult(success: true, message: json['message'] ?? 'Logged in');",
    '  } catch (_) {',
    '    return AuthResult(',
    '      success: false,',
    "      message: 'Network error. Check your internet connection',",
    '    );',
    '  }',
    '}',
])

add_heading('3. Active Backend Endpoints', 3)
add_table(
    ['Method', 'Endpoint', 'Purpose'],
    [
        ['POST', '/api/auth/signup',       'Register a new user'],
        ['GET',  '/api/auth/verify/:token','Verify email with token from link'],
        ['POST', '/api/auth/login',        'Authenticate and receive JWT'],
    ]
)

add_heading('4. Firebase Auth Integration', 3)
add_para('Firebase Auth is integrated as a secondary auth layer for password-management operations. During login the app also signs into Firebase so that password-change operations can be performed later:')
add_code([
    'Future<void> _signIntoFirebase(String email, String password) async {',
    '  try {',
    '    await _fb.signInWithEmailAndPassword(email: email, password: password);',
    '  } on FirebaseAuthException catch (e) {',
    "    if (e.code == 'user-not-found') {",
    '      try {',
    '        await _fb.createUserWithEmailAndPassword(email: email, password: password);',
    '      } catch (_) {}',
    '    }',
    '  }',
    '}',
])

add_heading('5. Translation Service — Multipart File Upload', 3)
add_para('The sign-language translation feature sends a video file to the AI server using http.MultipartRequest:')
add_code([
    'Future<PredictionResult> predict(File videoFile) async {',
    '  final request = http.MultipartRequest(',
    "    'POST',",
    "    Uri.parse('$_base/predict'),",
    '  );',
    '  request.files.add(',
    "    await http.MultipartFile.fromPath('video', videoFile.path),",
    '  );',
    '  final streamed = await request.send()',
    '      .timeout(const Duration(seconds: 60));',
    '  final res = await http.Response.fromStream(streamed);',
    '  if (res.statusCode != 200) throw Exception(\'Prediction failed\');',
    '  return PredictionResult.fromJson(jsonDecode(res.body));',
    '}',
])

add_heading('6. Chat Service — Session-Based Messaging', 3)
add_code([
    'class ChatService {',
    "  final _base = 'http://91.108.113.135';",
    '',
    '  Future<String> newSession() async {',
    '    final res = await http.post(',
    "      Uri.parse('$_base/session/new'),",
    "      headers: {'Content-Type': 'application/json'},",
    '    ).timeout(const Duration(seconds: 15));',
    "    return jsonDecode(res.body)['session_id'];",
    '  }',
    '',
    '  Future<String> sendMessage({',
    '    required String sessionId,',
    '    required String message,',
    '  }) async {',
    '    final res = await http.post(',
    "      Uri.parse('$_base/chat'),",
    "      headers: {'Content-Type': 'application/json'},",
    "      body: jsonEncode({'session_id': sessionId, 'message': message}),",
    '    ).timeout(const Duration(seconds: 30));',
    "    return jsonDecode(res.body)['response'];",
    '  }',
    '}',
])

add_heading('7. Error Handling and Timeouts', 3)
add_para('All API calls use:')
add_bullet('**.timeout(Duration)** — to prevent indefinite waits on slow connections (20 seconds for auth, up to 60 seconds for video upload)')
add_bullet('**try/catch blocks** — to catch network errors, JSON decode failures, and timeouts')
add_bullet('**Uniform AuthResult / Exception return** — so UI code always knows how to handle any outcome')

add_heading('8. Summary', 3)
add_para('The I Hear You application integrates three external services through a clean service-layer pattern using the standard http package. By encapsulating all HTTP logic in AuthService, TranslationService, and ChatService, the UI remains clean and free of network concerns. Consistent error handling, response parsing, and timeout management ensure a robust user experience even on unreliable connections.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7.5  Deployment and Device Testing
# ══════════════════════════════════════════════════════════════════════════════
add_heading('7.5  Deployment and Device Testing in Flutter', 2)

add_heading('1. Introduction', 3)
add_para('Developing a Flutter application is only part of the journey. Before users can benefit from the app, it must go through testing on real devices and eventually be deployed to official app stores such as Google Play and the Apple App Store. This phase is crucial for delivering a polished, functional, and reliable product.')

add_heading('2. Device Testing in Flutter', 3)

add_heading('2.1  Why Device Testing Matters', 4)
add_para('Testing on physical devices ensures:')
add_bullet('Real-world performance (battery, network, CPU usage)')
add_bullet('Accurate UI rendering on actual screen sizes and notches')
add_bullet('Access to hardware features (camera, microphone, storage) used in the app')
add_bullet('Correct behaviour of the speech_to_text microphone input in the Chat Screen')
add_bullet('Accurate performance of video_player and image_picker on real devices')

add_heading('2.2  Running on Android', 4)
add_bullet('Enable Developer Options and USB Debugging on the Android device')
add_bullet('Connect the device via USB')
add_bullet('Run: flutter run')

add_heading('2.3  Running on iOS', 4)
add_bullet('Open Xcode → Preferences → Accounts and sign in with an Apple ID')
add_bullet('Connect your iPhone via USB')
add_bullet('Run: flutter run')
add_para('Note: A macOS system is required to build and test iOS apps.')

add_heading('2.4  Using Emulators / Simulators', 4)
add_bullet('**Android Emulator:** Open Android Studio → AVD Manager → Create Virtual Device → Choose Pixel device with appropriate API level → Run flutter run')
add_bullet('**iOS Simulator:** Only available on macOS → Open via Xcode or run: open -a Simulator')

add_heading('3. Preparing the App for Release', 3)

add_heading('3.1  App Name, Icon, and Splash Screen', 4)
add_code([
    '# flutter_native_splash.yaml',
    'flutter_native_splash:',
    '  color: "#0B1020"',
    '  image: assets/images/splash_logo.png',
    '  android: true',
    '  ios: true',
    '',
    '# Terminal commands:',
    'flutter pub run flutter_launcher_icons:main',
    'flutter pub run flutter_native_splash:create',
])

add_heading('3.2  Required Permissions', 4)
add_table(
    ['Permission', 'Used By'],
    [
        ['Camera',         'Sign Practice Screen — video recording'],
        ['Microphone',     'AI Chat Screen — speech-to-text input'],
        ['Storage (read)', 'Image picker for gallery video selection'],
    ]
)

add_heading('4. Building the Release Version', 3)

add_heading('4.1  Android Release', 4)
add_code([
    '# Step 1: Create keystore',
    'keytool -genkey -v -keystore ~/my-release-key.jks \\',
    '  -keyalg RSA -keysize 2048 -validity 10000 -alias my-key-alias',
    '',
    '# Step 2: Build App Bundle for Play Store',
    'flutter build appbundle --release',
    '',
    '# Or APK for direct installation',
    'flutter build apk --release',
])

add_heading('4.2  iOS Release', 4)
add_bullet('Open project in Xcode: open ios/Runner.xcworkspace')
add_bullet('Update version and build number in Runner → General → Identity')
add_bullet('Configure Signing & Capabilities — select Apple Developer Team and unique Bundle Identifier')
add_bullet('Select Any iOS Device → Product → Archive')
add_bullet('Open Organizer to upload to the App Store')

add_heading('5. Store Deployment', 3)

add_heading('5.1  Google Play Store', 4)
add_bullet('Sign in to Google Play Console')
add_bullet('Create a new application')
add_bullet('Upload the App Bundle (.aab) file')
add_bullet('Fill in app name, description, screenshots, privacy policy, content rating')
add_bullet('Submit for review')

add_heading('5.2  Apple App Store', 4)
add_bullet('Sign in to App Store Connect')
add_bullet('Create a new app record')
add_bullet('Upload build via Xcode or Transporter')
add_bullet('Fill out app information, screenshots, pricing, and privacy compliance')
add_bullet('Submit for review')

add_heading('6. Testing Scenarios', 3)
add_para('The following scenarios were tested manually on the Android emulator:')
add_bullet('Complete authentication flow (signup → verify → login → logout)')
add_bullet('Full onboarding quiz (language → about you → purpose → level → study time)')
add_bullet('Flashcard learning flow (word list → I Know / Learn → sign practice)')
add_bullet('Translation screen (upload video → server health check → result display)')
add_bullet('AI chat (voice input → session management → response display)')
add_bullet('Profile settings (dark mode toggle → name edit → theme persistence)')
add_bullet('Theme switching in all screens (no flash of wrong theme)')
add_bullet('Language switching (all 14 languages render correctly)')

add_heading('7. Summary', 3)
add_para('Deploying the I Hear You Flutter app requires careful testing across screen sizes and OS versions, proper signing for both Android and iOS, and adherence to platform store guidelines. The use of flutter_native_splash, flutter_launcher_icons, and permission_handler ensures a polished, production-ready experience from the very first screen the user sees.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7.6  Sample Screenshots and Code Snippets
# ══════════════════════════════════════════════════════════════════════════════
add_heading('7.6  Sample Screenshots and Code Snippets in Flutter', 2)

add_heading('1. Introduction', 3)
add_para('Sample screenshots and code snippets are a vital part of technical documentation. They provide visual context and practical implementation examples that guide understanding of how the application works in practice. This section presents key screens of the I Hear You application together with the code structures that drive them.')

add_heading('2. Authentication Screens', 3)

add_heading('2.1  Splash Screen — Theme-Aware Entry Point', 4)
add_para('The splash screen loads the user\'s saved theme before rendering, determines the correct destination screen, and navigates after a short delay:')
add_code([
    '// splash_screen.dart',
    '@override',
    'void initState() {',
    '  super.initState();',
    '  _navigate();',
    '}',
    '',
    'Future<void> _navigate() async {',
    '  await Future.delayed(const Duration(seconds: 2));',
    '  final prefs = await SharedPreferences.getInstance();',
    "  final isLoggedIn       = prefs.getBool('isLoggedIn') ?? false;",
    "  final hasSelectedLevel = prefs.getBool('hasSelectedLevel') ?? false;",
    '  if (!mounted) return;',
    '',
    '  if (isLoggedIn && hasSelectedLevel) {',
    '    Navigator.pushReplacementNamed(context, Routes.dashboard);',
    '  } else if (isLoggedIn) {',
    '    Navigator.pushReplacementNamed(context, Routes.chooseLanguage);',
    '  } else {',
    '    Navigator.pushReplacementNamed(context, Routes.startPage);',
    '  }',
    '}',
])

add_heading('2.2  Login Screen — Form with Validation', 4)
add_code([
    '// login_screen.dart',
    'Future<void> _login() async {',
    '  if (!_formKey.currentState!.validate()) return;',
    '  setState(() => _isLoading = true);',
    '',
    '  final result = await AuthService().login(',
    '    email: _emailCtrl.text.trim(),',
    '    password: _passCtrl.text.trim(),',
    '  );',
    '',
    '  if (!mounted) return;',
    '  setState(() => _isLoading = false);',
    '',
    '  if (result.success || result.message.toLowerCase().contains(\'verify\')) {',
    '    final prefs = await SharedPreferences.getInstance();',
    "    await prefs.setBool('isLoggedIn', true);",
    "    final hasLevel = prefs.getBool('hasSelectedLevel') ?? false;",
    '    Navigator.pushNamedAndRemoveUntil(',
    '      context,',
    '      hasLevel ? Routes.dashboard : Routes.chooseLanguage,',
    '      (_) => false,',
    '    );',
    '  } else {',
    '    ScaffoldMessenger.of(context).showSnackBar(',
    '      SnackBar(content: Text(result.message), backgroundColor: Colors.red),',
    '    );',
    '  }',
    '}',
])

add_heading('3. Onboarding Quiz Screens', 3)

add_heading('3.1  Level Options Screen — CEFR Level Selection', 4)
add_code([
    '// level_options_screen.dart',
    'void _selectLevel(String level) async {',
    '  setState(() => _selectedLevel = level);',
    '  final prefs = await SharedPreferences.getInstance();',
    "  await prefs.setString('userLevel', level);",
    '}',
    '',
    "Widget _buildLevelCard(String code, String label, String description) {",
    '  final isSelected = _selectedLevel == code;',
    '  return GestureDetector(',
    '    onTap: () => _selectLevel(code),',
    '    child: Container(',
    '      padding: EdgeInsets.all(16.w),',
    '      decoration: BoxDecoration(',
    '        color: isSelected ? accent.withValues(alpha: 0.15) : cardColor,',
    '        border: Border.all(',
    '          color: isSelected ? accent : Colors.transparent, width: 2),',
    '        borderRadius: BorderRadius.circular(16.r),',
    '      ),',
    '      child: Row( /* level code + label + description + check icon */ ),',
    '    ),',
    '  );',
    '}',
])

add_heading('4. Learning Screens', 3)

add_heading('4.1  Word List Screen — Flashcard Interface', 4)
add_code([
    '// word_list_screen.dart',
    'void _markAsKnown(String word) async {',
    '  final prefs = await SharedPreferences.getInstance();',
    "  final knownWords = prefs.getStringList('known_$_level') ?? [];",
    '  if (!knownWords.contains(word)) {',
    '    knownWords.add(word);',
    "    await prefs.setStringList('known_$_level', knownWords);",
    '  }',
    '  setState(() => _currentIndex++);',
    '  if (_currentIndex >= _words.length) {',
    '    _navigateToSignPractice();',
    '  }',
    '}',
    '',
    'void _markForLearning() {',
    '  setState(() {',
    '    final word = _words.removeAt(_currentIndex);',
    '    _words.add(word);  // move to end of queue',
    '  });',
    '}',
])

add_heading('4.2  Sign Practice Screen — Video Analysis', 4)
add_code([
    '// sign_practice_screen.dart',
    'Future<void> _pickAndAnalyze() async {',
    '  final picker = ImagePicker();',
    '  final picked = await picker.pickVideo(source: ImageSource.gallery);',
    '  if (picked == null) return;',
    '  final file = File(picked.path);',
    '  setState(() => _selectedVideo = file);',
    '  final notifier = context.read<TranslationNotifier>();',
    '  await notifier.predict(file);',
    '}',
    '',
    '// Reactive result display:',
    'if (notifier.isLoading) return const CircularProgressIndicator();',
    'if (notifier.isSuccess) {',
    '  return ResultCard(',
    '    predicted: notifier.result!.top.label,',
    '    expected: widget.wordToSign,',
    '  );',
    '}',
])

add_heading('5. Translation Screen (Lost Page Screen)', 3)
add_code([
    '// lost_page_screen.dart',
    '@override',
    'void initState() {',
    '  super.initState();',
    '  _notifier = TranslationNotifier();',
    '  _notifier.checkHealth();',
    '  _notifier.loadClasses();',
    '  _notifier.addListener(_rebuild);',
    '}',
    '',
    'Widget _buildBody() {',
    '  if (_notifier.isLoading)',
    '    return const Center(child: CircularProgressIndicator());',
    '  if (_notifier.isSuccess)',
    '    return ResultView(result: _notifier.result!, sentence: _notifier.sentence);',
    '  if (_notifier.isError)',
    '    return ErrorView(message: _notifier.error, onRetry: _notifier.reset);',
    '  return UploadPromptView(onPick: _pickVideo);',
    '}',
])

add_heading('6. AI Chat Screen', 3)
add_code([
    '// chat_screen.dart',
    'Future<void> _sendMessage([String? override]) async {',
    '  final text = override ?? _inputController.text.trim();',
    '  if (text.isEmpty) return;',
    '',
    '  setState(() {',
    '    _messages.add(_ChatMessage(text: text, isUser: true));',
    '    _isTyping = true;',
    '  });',
    '  _inputController.clear();',
    '',
    '  String reply;',
    '  try {',
    '    final sessionId = await _ensureSession();',
    '    reply = await _service.sendMessage(',
    '      sessionId: sessionId, message: text);',
    '  } catch (_) {',
    '    reply = "Sorry, I could not reach the server.";',
    '  }',
    '',
    '  setState(() {',
    '    _isTyping = false;',
    '    _messages.add(_ChatMessage(text: reply, isUser: false));',
    '  });',
    '}',
])

add_heading('7. Profile Screen — Settings Hub', 3)
add_code([
    '// profile._screen.dart',
    'profileTile(',
    '  icon: Icons.dark_mode_outlined,',
    "  title: S.get('dark_mode'),",
    '  trailing: Transform.scale(',
    '    scale: 0.85,',
    '    child: Switch(',
    '      value: isDark,',
    '      activeTrackColor: accent,',
    '      onChanged: (value) {',
    '        setState(() => isDark = value);',
    '        themeNotifier.value = value ? ThemeMode.dark : ThemeMode.light;',
    '        prefs.setBool(\'isDark\', value);',
    '      },',
    '    ),',
    '  ),',
    '),',
    '',
    'profileTile(',
    '  icon: Icons.lock_outline_rounded,',
    "  title: S.get('password'),",
    '  onTap: () => ScaffoldMessenger.of(context).showSnackBar(',
    '    SnackBar(',
    "      content: Text(S.get('coming_soon')),",
    '      backgroundColor: const Color(0xFF22C55E),',
    '    ),',
    '  ),',
    '),',
])

add_heading('8. Multi-Language Support', 3)
add_para('All user-facing strings are served through the S.get(key) helper, which looks up the current language from languageNotifier and falls back to English if a translation is missing:')
add_code([
    '// app_strings.dart',
    'class S {',
    '  static String get(String key) {',
    '    final code = languageNotifier.value.code;',
    "    return _t[code]?[key] ?? _t['en']![key] ?? key;",
    '  }',
    '',
    '  static const _t = <String, Map<String, String>>{',
    "    'en': {",
    "      'login': 'Login',",
    "      'signup': 'Sign Up',",
    "      'get_started': 'Get Started',",
    "      'dark_mode': 'Dark Mode',",
    "      'coming_soon': 'Coming Soon',",
    '      // ... 80+ keys',
    '    },',
    "    'ar': {",
    "      'login': 'تسجيل الدخول',",
    "      'signup': 'إنشاء حساب',",
    "      'get_started': 'ابدأ الآن',",
    "      'dark_mode': 'الوضع الداكن',",
    "      'coming_soon': 'قريباً',",
    '    },',
    '    // + 12 more: sp, fr, ru, ge, it, tu, ja, ch, ko, po, hi, du',
    '  };',
    '}',
])

add_heading('9. Key Packages and Dependencies', 3)
add_table(
    ['Package', 'Version', 'Purpose'],
    [
        ['firebase_auth',             '^6.5.2',  'Password management and secondary auth'],
        ['firebase_core',             '^4.10.0', 'Firebase initialisation'],
        ['cloud_firestore',           '^6.5.0',  'Cloud database (available for future features)'],
        ['flutter_screenutil',        '^5.9.3',  'Responsive UI scaling (designSize 390×844)'],
        ['http',                      '^1.2.1',  'REST API calls to backend and AI servers'],
        ['shared_preferences',        '^2.2.3',  'Persistent local storage for user state'],
        ['flutter_svg',               '^2.2.4',  'SVG asset rendering (app logo)'],
        ['google_fonts',              '^8.0.2',  'Poppins, Open Sans, Inter font families'],
        ['video_player',              '^2.9.1',  'Video preview in sign practice screen'],
        ['image_picker',              '^1.1.2',  'Gallery and camera video selection'],
        ['speech_to_text',            '^7.0.0',  'Voice input for AI chat screen'],
        ['permission_handler',        '^11.3.1', 'Runtime permission requests'],
        ['flutter_local_notifications','^17.2.2','Local push notifications'],
        ['url_launcher',              '^6.3.1',  'Opening external links'],
        ['flutter_native_splash',     '^2.4.7',  'Native splash screen configuration'],
        ['flutter_launcher_icons',    '^0.14.4', 'App icon generation for Android and iOS'],
    ]
)

add_heading('10. Summary', 3)
add_para('The I Hear You Flutter application demonstrates a well-structured, feature-complete mobile app built with modern Flutter practices. The combination of a feature-based folder structure, lightweight state management (ValueNotifier, ChangeNotifier, SharedPreferences), dual-backend integration (REST + Firebase), full responsive design via flutter_screenutil, 14-language support, and async safety patterns produces a polished, accessible sign-language learning experience across both Android and iOS platforms.')

# ── Save ──────────────────────────────────────────────────────────────────────
out = r"c:\Users\Osha\Downloads\popsign-main (1)\popsign-main\Flutter_Chapter.docx"
doc.save(out)
print(f"Saved: {out}")
